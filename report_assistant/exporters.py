"""把 Markdown 报告导出为 .md / .docx / .pdf。

Markdown 解析采用轻量自实现（不引入 markdown 库），覆盖：
- # / ## / ### 标题
- 无序列表 (- / *) 与有序列表 (1.)
- > blockquote
- ``` 代码块（围栏）
- 段落内 **粗体** 和 `行内代码`
- 水平线 ---

足以渲染我们 LLM 生成的报告（生成 prompt 已限定输出风格）。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional


# ── Markdown 解析（最小够用） ─────────────────────

_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_FENCE = re.compile(r"^```(\w*)\s*$")
_OL_ITEM = re.compile(r"^\s*\d+\.\s+(.+?)\s*$")
_UL_ITEM = re.compile(r"^\s*[-*+]\s+(.+?)\s*$")
_BQUOTE = re.compile(r"^>\s?(.*)$")
_HR = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")

_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_CODE = re.compile(r"`([^`]+?)`")


def _parse_blocks(md: str) -> list[dict]:
    """把 markdown 切成块列表。"""
    lines = md.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        m = _FENCE.match(line)
        if m:
            lang = m.group(1) or ""
            buf = []
            i += 1
            while i < len(lines) and not _FENCE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(buf)})
            continue

        if not line.strip():
            i += 1
            continue

        if _HR.match(line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        m = _HEADING.match(line)
        if m:
            blocks.append({
                "type": "heading",
                "level": len(m.group(1)),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        if _BQUOTE.match(line):
            buf = []
            while i < len(lines):
                qm = _BQUOTE.match(lines[i])
                if not qm:
                    break
                buf.append(qm.group(1))
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(buf).strip()})
            continue

        if _UL_ITEM.match(line):
            items = []
            while i < len(lines):
                um = _UL_ITEM.match(lines[i])
                if not um:
                    break
                items.append(um.group(1).strip())
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue

        if _OL_ITEM.match(line):
            items = []
            while i < len(lines):
                om = _OL_ITEM.match(lines[i])
                if not om:
                    break
                items.append(om.group(1).strip())
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue

        buf = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (not nxt.strip()
                    or _HEADING.match(nxt) or _FENCE.match(nxt)
                    or _BQUOTE.match(nxt) or _UL_ITEM.match(nxt)
                    or _OL_ITEM.match(nxt) or _HR.match(nxt)):
                break
            buf.append(nxt)
            i += 1
        blocks.append({"type": "para", "text": " ".join(s.strip() for s in buf)})
    return blocks


def _split_inline(text: str) -> list[tuple[str, str]]:
    """把行内文本拆成 [(kind, text), ...]，kind ∈ bold / code / plain。"""
    segments: list[tuple[str, str]] = [("plain", text)]

    def _split_by(segs, regex, kind):
        out = []
        for k, s in segs:
            if k != "plain":
                out.append((k, s))
                continue
            last = 0
            for m in regex.finditer(s):
                if m.start() > last:
                    out.append(("plain", s[last:m.start()]))
                out.append((kind, m.group(1)))
                last = m.end()
            if last < len(s):
                out.append(("plain", s[last:]))
        return out

    segments = _split_by(segments, _INLINE_BOLD, "bold")
    segments = _split_by(segments, _INLINE_CODE, "code")
    return [(k, v) for k, v in segments if v]


# ── 导出：Markdown ────────────────────────────────

def export_markdown(content: str, target: str | Path) -> Path:
    p = Path(str(target))
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
    return p


# ── 导出：DOCX ────────────────────────────────────

def export_docx(content: str, target: str | Path) -> Path:
    """用 python-docx 渲染 markdown 报告为 .docx。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
    except ImportError as e:
        raise RuntimeError("请先安装 python-docx：pip install python-docx") from e

    blocks = _parse_blocks(content)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    def _emit_runs(p, text):
        for kind, txt in _split_inline(text):
            run = p.add_run(txt)
            if kind == "bold":
                run.bold = True
            elif kind == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)

    for blk in blocks:
        bt = blk["type"]
        if bt == "heading":
            level = min(blk["level"], 4)
            p = doc.add_heading(blk["text"], level=level)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
        elif bt == "para":
            p = doc.add_paragraph()
            _emit_runs(p, blk["text"])
        elif bt == "ul":
            for item in blk["items"]:
                p = doc.add_paragraph(style="List Bullet")
                _emit_runs(p, item)
        elif bt == "ol":
            for item in blk["items"]:
                p = doc.add_paragraph(style="List Number")
                _emit_runs(p, item)
        elif bt == "quote":
            p = doc.add_paragraph(blk["text"])
            p.paragraph_format.left_indent = Cm(0.6)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
        elif bt == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            run = p.add_run(blk["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif bt == "hr":
            doc.add_paragraph("─" * 40)

    target_path = Path(str(target))
    target_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target_path))
    return target_path


# ── 导出：PDF（需要 PySide6） ─────────────────────

_PDF_CSS = """
body { font-family: "Microsoft YaHei UI", "PingFang SC", sans-serif;
       color: #0f172a; line-height: 1.7; font-size: 11pt; }
h1 { font-size: 20pt; color: #0f172a; border-bottom: 1px solid #e2e8f0;
     padding-bottom: 6px; margin: 14px 0 10px; }
h2 { font-size: 15pt; color: #047857; margin: 14px 0 8px; }
h3 { font-size: 13pt; color: #334155; margin: 10px 0 6px; }
p  { margin: 5px 0; color: #334155; }
ul, ol { margin: 6px 0 6px 18px; color: #334155; }
li { margin: 3px 0; }
strong, b { color: #0f172a; font-weight: 600; }
code { background: #f1f5f9; color: #0f172a; padding: 1px 5px;
       border-radius: 3px; font-family: "Consolas", monospace; font-size: 10pt; }
pre  { background: #f8fafc; border: 1px solid #e2e8f0;
       border-radius: 6px; padding: 10px; }
pre code { background: transparent; padding: 0; }
blockquote { border-left: 3px solid #10b981; background: #ecfdf5;
             padding: 6px 12px; margin: 6px 0; color: #047857; }
hr { border: 0; height: 1px; background: #e2e8f0; margin: 12px 0; }
"""


def export_pdf(content: str, target: str | Path) -> Path:
    """用 QTextDocument + QPdfWriter 渲染。需要 PySide6 + 已有 QApplication。"""
    try:
        from PySide6.QtCore import QMarginsF
        from PySide6.QtGui import QPageLayout, QPageSize, QPdfWriter, QTextDocument
        from PySide6.QtWidgets import QApplication
    except ImportError as e:
        raise RuntimeError("请先安装 PySide6") from e

    if QApplication.instance() is None:
        import sys
        QApplication(sys.argv)

    target_path = Path(str(target))
    target_path.parent.mkdir(parents=True, exist_ok=True)

    doc = QTextDocument()
    doc.setDefaultStyleSheet(_PDF_CSS)
    doc.setMarkdown(content)

    writer = QPdfWriter(str(target_path))
    writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
    layout = writer.pageLayout()
    layout.setUnits(QPageLayout.Unit.Millimeter)
    layout.setMargins(QMarginsF(20, 18, 20, 18))
    writer.setPageLayout(layout)
    writer.setResolution(150)

    doc.setPageSize(
        writer.pageLayout().paintRectPixels(writer.resolution()).size().toSizeF()
    )
    doc.print_(writer)
    return target_path


# ── 统一入口 ──────────────────────────────────────

EXPORT_FORMATS: dict[str, str] = {
    "md": "Markdown (*.md)",
    "docx": "Word 文档 (*.docx)",
    "pdf": "PDF 文档 (*.pdf)",
}


def export_report(content: str, target: str | Path, fmt: Optional[str] = None) -> Path:
    """根据 target 的扩展名（或显式 fmt）导出。"""
    p = Path(str(target))
    ext = (fmt or p.suffix.lstrip(".")).lower()
    if ext == "md":
        return export_markdown(content, p)
    if ext == "docx":
        return export_docx(content, p)
    if ext == "pdf":
        return export_pdf(content, p)
    raise ValueError(f"不支持的导出格式: {ext}（支持 md/docx/pdf）")


def build_save_filter() -> str:
    return ";;".join(EXPORT_FORMATS.values())
